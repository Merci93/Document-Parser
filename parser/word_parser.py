"""
Extract the contents of a Word document, including table of contents, text, images and tables.
Note that .doc files from Word 2003 and earlier will not work
"""
import csv
import os
import re

import docx
import docx2txt
import pandas
from docx import Document
from docx.image.exceptions import UnrecognizedImageError
from docx.opc.exceptions import PackageNotFoundError
from unidecode import unidecode
from xml.etree import ElementTree as ET
from zipfile import BadZipFile

from configuration import settings
from document_parser import DocumentParser
from logger import log


class WordParser(DocumentParser):
    """A class responsible for reading and parsing word documents."""

    def __init__(self, document_path: str, output_directory: str = None) -> None:
        """
        Open Word document.

        As for different purposes two independent parsing libraries are used: docx (python-docx) and docx2txt,
        both parsers are loaded here.

        :param document_path: Full path to location of documents.
        :param output_directory: Location where parsed data are saved. If not provided, execution will not succeed.
        """
        super().__init__(document_path, output_directory)
        self.document_docx = Document(document_path)
        self.document_docx2txt = docx2txt.process(document_path)

    def extract_table_of_content(self) -> None:
        """Extract table of content from Word document."""
        final_directory = os.path.join(self.output_directory, settings.extracted_table_of_content)
        os.makedirs(final_directory, exist_ok=True)
        table_of_content = []
        for line in self.document_docx2txt.split("\n"):
            line = line.strip()
            match_pattern = re.match(r"\d+(\.\d+)*\s+[A-Za-z].*\d+$", line)
            if match_pattern:
                try:
                    text = line.rsplit("\t", 1)[0]
                    page_number = line.rsplit("\t", 1)[1]
                    table_of_content.append((text, page_number))
                except IndexError:
                    pass
        log.info(f"Extracted TOC with {len(table_of_content)} titles.")
        if len(table_of_content) > 0:
            toc_df = pandas.DataFrame(table_of_content, columns=["title", "page_number"])
            toc_df.to_csv(os.path.join(final_directory, f"{self.document_name}.csv"), index=False)
        self.counters['toc'] = len(table_of_content)

    def extract_images(self) -> None:
        """Extract images with figure labels from document and save them."""
        log.info("Extracting images...")
        final_directory = os.path.join(self.output_directory, settings.extracted_images)
        os.makedirs(final_directory, exist_ok=True)

        def found_image_with_title(paragraph: docx.text.paragraph) -> docx.ImagePart | bool:
            """Extract image blob data and rId from paragraph xml data."""
            for run in paragraph.runs:
                for inline in run._r.xpath("w:drawing/wp:inline"):
                    image_rId = inline.graphic.graphicData.pic.blipFill.blip.embed
                    try:
                        image = self.document_docx.part.related_parts[image_rId].image
                        return image
                    except UnrecognizedImageError:
                        return False

        def save_image(figure_match: re, data: dict, image: docx.ImagePart, counter: int) -> dict:
            """Save extracted image."""
            save_as = f"{self.document_name}_{figure_match.group(1)}{counter}.png"
            if image:
                with open(os.path.join(final_directory, save_as), "wb") as img:
                    img.write(image.blob)
                data.update(
                    {
                        "figure_number": f"{figure_match.group(1)}{counter}",
                        "figure_title": figure_match.group(2), "image_filename": save_as,
                    }
                )
                return data
            else:
                data.update(
                    {
                        "figure_number": f"{figure_match.group(1)}{counter}",
                        "figure_title": figure_match.group(2), "image_filename": "Image not extracted",
                        "extracted_image": "No. Image has a bad format.",
                    }
                )
                return data

        extracted_images = []
        image_titles = []
        paragraph = self.document_docx.paragraphs
        for i in range(len(paragraph)):
            if "graphicData" in paragraph[i]._p.xml:
                prev_paragraph = paragraph[i - 1]
                next_paragraph = paragraph[i + 1]
                if prev_paragraph.text.startswith("Figure ") and "Caption" in prev_paragraph.style.name:
                    image_titles.append(prev_paragraph.text)
                    extracted_images.append(found_image_with_title(paragraph[i]))

                elif prev_paragraph.text == "":
                    for j in range(2, 4):
                        prev_paragraph = paragraph[i - j]
                        if prev_paragraph.text.startswith("Figure ") and "Caption" in prev_paragraph.style.name:
                            image_titles.append(prev_paragraph.text)
                            extracted_images.append(found_image_with_title(paragraph[i]))
                            break

                elif next_paragraph.text.startswith("Figure ") and "Caption" in next_paragraph.style.name:
                    image_titles.append(next_paragraph.text)
                    extracted_images.append(found_image_with_title(paragraph[i]))

        counter = 0
        image_data = []
        zipped_file = dict(zip(image_titles, extracted_images))
        for title, image in zipped_file.items():
            counter += 1
            title = unidecode(title.strip())
            figure_title = re.match(r"^(Figure\s+\d+(?:[:\.-]\d+)?)\s*(.*)", title)
            figure_title_part = re.match(r"^(Figure\s+)(?:[\s:-])-?(.*)", title)
            data = {
                "document_name": self.document_name,
                "document_type": "Word",
                "figure_number": "figure number",
                "figure_title": "figure title",
                "page_number": "N/A",
                "image_filename": "save_as",
                "extracted_image": "Yes"
            }
            if figure_title:
                image_data.append(save_image(figure_title, data, image, ""))
            elif figure_title_part:
                if counter == 1:
                    image_data.append(save_image(figure_title_part, data, image, counter))
                elif counter > 1:
                    number = int(image_data[-1]["figure_number"].rsplit(" ", 1)[1]) + 1
                    image_data.append(save_image(figure_title_part, data, image, number))
        log.info(f"...done. Successfully extracted {counter} images.")
        self.counters["images"] = counter
        self.image_data = image_data

    def extract_tables(self) -> None:
        """Extract tables from Word document and save them."""
        log.info("Extracting tables...")
        final_directory = os.path.join(self.output_directory, settings.extracted_tables)
        os.makedirs(final_directory, exist_ok=True)

        table_titles = []
        name_space = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        for idx, element in enumerate(self.document_docx.element.body):
            if (element.tag.endswith("tbl")) and (idx > 0):
                title = ET.fromstring(self.document_docx.element.body[idx - 1].xml)
                tags = []
                for wt_tag in title.findall(".//w:r//w:t", name_space):
                    tags.append(wt_tag.text)
                table_title = "".join(tags)
                table_titles.append(table_title)

        tables = self.document_docx.tables
        extracted_tables = []
        for idx, table in enumerate(tables, start=1):
            extracted_table = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                extracted_table.append(row_text)
            extracted_tables.append(extracted_table)

        zip_table_title = dict(zip(table_titles, extracted_tables))
        counter = 0
        for title, table in zip_table_title.items():
            table_title = re.match(r"^Table\s+\d+", title.strip())
            if (title.replace(" ", "") != "") and table_title:
                counter += 1
                df = pandas.DataFrame(table)
                df = df.rename(columns=df.iloc[0]).loc[1:]
                headers = [f"{title}"]
                if df.shape[1] != 1:
                    headers = headers + ([""] * (df.shape[1] - 1))
                df.columns = pandas.MultiIndex.from_tuples(zip(headers, df.columns))
                save_as = f"{self.document_name}_{counter}.csv"
                df.to_csv(os.path.join(final_directory, save_as), index=False)
        log.info(f"...done. Successfully extracted {counter} tables.")
        self.counters['tables'] = counter

    def extract_texts(self) -> None:
        """Extract text from Word document and save as CSV file."""
        log.info("Extracting text...")
        final_directory = os.path.join(self.output_directory, settings.extracted_texts)
        os.makedirs(final_directory, exist_ok=True)

        extracted_sentences = []
        for paragraph in self.document_docx.paragraphs:
            text = paragraph.text
            table_title = re.match(r"^Table\s+\d+", text)
            figure_title = re.match(r"^Figure\s+\d+", text)
            if (text.replace(" ", "") != "") and (not table_title) and (not figure_title):
                extracted_sentences.append((text, paragraph.style.name))
        sentence_df = pandas.DataFrame(extracted_sentences, columns=["text", "text_style"])

        headers = []
        body_text = []
        tmp = []
        for _, row_value in sentence_df.iterrows():
            text = unidecode(row_value.text)
            text_style = row_value.text_style
            if text_style.startswith("Heading"):
                headers.append(text)
                body_text.append("\n".join(tmp))
                tmp = []
            else:
                tmp.append(text)
        body_text.append("\n".join(tmp))
        body_text = body_text[1:]
        text_df = pandas.DataFrame(zip(headers, body_text), columns=["heading", "text"])
        text_df.to_csv(os.path.join(final_directory, f"{self.document_name}.csv"), index=False)
        log.info(f"...done. Successfully extracted {len(headers)} paragraphs.")
        self.counters['paragraphs'] = len(headers)


def parse_all_word_documents(
    input_directory: str = settings.file_location,
    output_directory: str = settings.parsed_data_directory,
    ) -> None:
    """Parse all pdf documents that exist in input_directory and save results to output_directory."""
    image_counter = []
    table_counter = []
    for document in os.listdir(input_directory):
        if document.endswith("docx"):
            try:
                parser = WordParser(os.path.join(input_directory, document), output_directory)
            except (PackageNotFoundError, BadZipFile) as e:
                log.error(f"{document} cannot be parsed. Error: {e}")
                break
            parser.extract_table_of_content()
            parser.extract_texts()
            table_counter.append((parser.document_name, parser.extract_tables()))
            image_counter.append((parser.document_name, parser.extract_images()))
            log.info(f"{document} successfully processed.")
        else:
            log.info(f"Skipping not a word document {document}.")
    with open(os.path.join(output_directory, "word_doc_image_report.csv"), "w", newline="") as image_report_file:
        csv.writer(image_report_file).writerows(image_counter)
    with open(os.path.join(output_directory, "word_doc_table_report.csv"), "w", newline="") as table_report_file:
        csv.writer(table_report_file).writerows(table_counter)


if __name__ == "__main__":
    parse_all_word_documents()
    